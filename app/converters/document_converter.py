from pathlib import Path
from zipfile import ZipFile
from xml.etree import ElementTree


def docx_to_txt(input_path, output_folder):
    input_path = Path(input_path)
    output_path = Path(output_folder) / f"{input_path.stem}.txt"
    namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    paragraphs = []

    with ZipFile(input_path) as archive:
        document_xml = archive.read("word/document.xml")

    root = ElementTree.fromstring(document_xml)
    for paragraph in root.iter(f"{namespace}p"):
        text = "".join(node.text or "" for node in paragraph.iter(f"{namespace}t"))
        paragraphs.append(text)

    output_path.write_text("\n".join(paragraphs), encoding="utf-8")
    return output_path


def txt_to_docx(input_path, output_folder):
    input_path = Path(input_path)
    output_path = Path(output_folder) / f"{input_path.stem}.docx"

    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError(
            "DOCX creation requires python-docx. Install it with: pip install python-docx"
        ) from error

    document = Document()
    for line in input_path.read_text(encoding="utf-8").splitlines():
        document.add_paragraph(line)
    document.save(output_path)
    return output_path


def text_to_pdf(input_path, output_folder):
    input_path = Path(input_path)
    output_path = Path(output_folder) / f"{input_path.stem}.pdf"

    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.pdfgen.canvas import Canvas
    except ImportError as error:
        raise RuntimeError(
            "PDF creation requires reportlab. Install it with: pip install reportlab"
        ) from error

    canvas = Canvas(str(output_path), pagesize=A4)
    width, height = A4
    y = height - 48
    canvas.setFont("Helvetica", 10)

    for line in input_path.read_text(encoding="utf-8").splitlines():
        canvas.drawString(48, y, line[:120])
        y -= 15
        if y < 48:
            canvas.showPage()
            canvas.setFont("Helvetica", 10)
            y = height - 48

    canvas.save()
    return output_path


def docx_to_pdf(input_path, output_folder):
    input_path = Path(input_path)
    text_path = docx_to_txt(input_path, output_folder)
    output_path = text_to_pdf(text_path, output_folder)
    text_path.unlink(missing_ok=True)
    return output_path


def pdf_to_docx(input_path, output_folder, cancel_event=None, progress_callback=None):
    input_path = Path(input_path)
    output_path = Path(output_folder) / f"{input_path.stem}.docx"

    try:
        import fitz
    except ImportError as error:
        raise RuntimeError(
            "PDF reading requires PyMuPDF. Install it with: pip install -r requirements.txt"
        ) from error

    try:
        from docx import Document
    except ImportError as error:
        raise RuntimeError(
            "DOCX creation requires python-docx. Install it with: pip install -r requirements.txt"
        ) from error

    document = Document()
    source = fitz.open(input_path)

    try:
        total_pages = len(source)
        for page_number, page in enumerate(source):
            if cancel_event and cancel_event.is_set():
                break
            if page_number:
                document.add_page_break()

            text = page.get_text("text").strip()
            for line in text.splitlines():
                document.add_paragraph(line)
            if progress_callback:
                progress_callback(page_number + 1, total_pages)
    finally:
        source.close()

    document.save(output_path)
    return output_path


def pdf_to_txt(input_path, output_folder, cancel_event=None, progress_callback=None):
    input_path = Path(input_path)
    output_path = Path(output_folder) / f"{input_path.stem}.txt"

    try:
        import fitz
    except ImportError as error:
        raise RuntimeError(
            "PDF reading requires PyMuPDF. Install it with: pip install -r requirements.txt"
        ) from error

    source = fitz.open(input_path)

    try:
        with output_path.open("w", encoding="utf-8") as output:
            total_pages = len(source)
            for page_number, page in enumerate(source, start=1):
                if cancel_event and cancel_event.is_set():
                    break
                if page_number > 1:
                    output.write("\n\n")
                output.write(page.get_text("text").strip())
                if progress_callback:
                    progress_callback(page_number, total_pages)
    finally:
        source.close()

    return output_path